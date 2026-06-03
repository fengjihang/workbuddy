"""文档解析：.docx / .pdf → 结构化文本"""

from pathlib import Path


class ParsedDocument:
    def __init__(self, full_text: str, sections: list[dict] | None = None):
        self.full_text = full_text
        self.sections = sections or []  # [{title, content, page}]


def parse_document(file_path: str) -> ParsedDocument:
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".pdf":
        return _parse_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _parse_docx(file_path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    sections: list[dict] = []
    current_section = {"title": "正文", "content": "", "page": 1}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        # 将加粗、字号较大的段落识别为章节标题
        is_heading = False
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            is_heading = True
        elif para.runs and para.runs[0].bold:
            is_heading = True

        if is_heading and current_section["content"]:
            sections.append(current_section)
            current_section = {"title": text, "content": "", "page": len(sections) + 1}
        else:
            current_section["content"] += text + "\n"

    if current_section["content"]:
        sections.append(current_section)

    # 提取表格内容
    for i, table in enumerate(doc.tables):
        table_text = _extract_table_text(table)
        if table_text:
            sections.append({"title": f"表格{i + 1}", "content": table_text, "page": len(sections) + 1})

    return ParsedDocument(full_text="\n".join(paragraphs), sections=sections)


def _parse_pdf(file_path: str) -> ParsedDocument:
    import fitz

    doc = fitz.open(file_path)
    full_text_parts = []
    sections: list[dict] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        full_text_parts.append(text)
        sections.append({
            "title": f"第{page_num + 1}页",
            "content": text,
            "page": page_num + 1,
        })

    doc.close()
    return ParsedDocument(full_text="\n".join(full_text_parts), sections=sections)


def _extract_table_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
