import json
import re
import os
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from ..database import get_db
from ..models.bid import Bid, BidChapter
from ..models.tender import Tender, AnalysisModule
from ..schemas.bid import BidCreate, BidOut, BidDetail, BidChapterOut, FieldListOut, FieldOccurrence, FieldInfo, InspectResult, MissingField
from ..services.bid_generator import generate_chapter, generate_chapter_stream

router = APIRouter(prefix="/api/bids", tags=["标书"])


@router.post("", response_model=BidOut)
def create_bid(body: BidCreate, db: Session = Depends(get_db)):
    bid = Bid(tender_id=body.tender_id, name=body.name)
    db.add(bid)
    db.commit()
    db.refresh(bid)

    # 初始化默认章节
    default_chapters = [
        "投标函", "法定代表人授权委托书", "投标人基本情况表",
        "技术方案", "项目实施方案", "售后服务方案",
        "商务方案", "报价表", "资质证明文件",
    ]
    for i, title in enumerate(default_chapters):
        db.add(BidChapter(bid_id=bid.id, chapter_index=i + 1, title=title))
    db.commit()

    return bid


@router.get("", response_model=list[BidOut])
def list_bids(db: Session = Depends(get_db)):
    return db.query(Bid).order_by(Bid.create_time.desc()).all()


@router.get("/{bid_id}/generate/{chapter_index}")
async def generate_chapter_content(bid_id: int, chapter_index: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    chapter = db.query(BidChapter).filter(
        BidChapter.bid_id == bid_id, BidChapter.chapter_index == chapter_index
    ).first()
    if not chapter:
        raise HTTPException(404, "章节不存在")

    # 收集解读上下文
    tender_context = ""
    if bid.tender_id:
        modules = db.query(AnalysisModule).filter(
            AnalysisModule.tender_id == bid.tender_id
        ).order_by(AnalysisModule.module_index).all()
        tender_context = "\n\n".join([
            f"## {m.module_name}\n{m.content}"
            for m in modules if m.content
        ])

    async def generate():
        async for msg in generate_chapter_stream(chapter.title, chapter_index, tender_context):
            yield f"data: {msg}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@router.get("/{bid_id}/fields", response_model=FieldListOut)
def extract_fields(bid_id: int, chapter_index: int | None = None, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    query = db.query(BidChapter).filter(BidChapter.bid_id == bid_id)
    if chapter_index is not None:
        query = query.filter(BidChapter.chapter_index == chapter_index)
    chapters = query.order_by(BidChapter.chapter_index).all()

    placeholder_pattern = re.compile(r'\{\{([^{}]+?)\}\}')
    field_map: dict[str, list[FieldOccurrence]] = {}

    for ch in chapters:
        if not ch.content:
            continue
        for match in placeholder_pattern.finditer(ch.content):
            field_name = match.group(1).strip()
            if field_name not in field_map:
                field_map[field_name] = []
            existing_indices = {o.chapter_index for o in field_map[field_name]}
            if ch.chapter_index not in existing_indices:
                field_map[field_name].append(
                    FieldOccurrence(chapter_index=ch.chapter_index, chapter_title=ch.title)
                )

    fields = [
        FieldInfo(field_name=name, occurrences=occs)
        for name, occs in sorted(field_map.items())
    ]

    return FieldListOut(bid_id=bid_id, fields=fields, total_count=len(fields))


@router.get("/{bid_id}/fields/excel")
def export_fields_excel(bid_id: int, chapter_index: int | None = None, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    query = db.query(BidChapter).filter(BidChapter.bid_id == bid_id)
    if chapter_index is not None:
        query = query.filter(BidChapter.chapter_index == chapter_index)
    chapters = query.order_by(BidChapter.chapter_index).all()

    placeholder_pattern = re.compile(r'\{\{([^{}]+?)\}\}')
    field_map: dict[str, list[str]] = {}
    for ch in chapters:
        if not ch.content:
            continue
        for match in placeholder_pattern.finditer(ch.content):
            field_name = match.group(1).strip()
            if field_name not in field_map:
                field_map[field_name] = []
            if ch.title not in field_map[field_name]:
                field_map[field_name].append(ch.title)

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, Protection

    wb = Workbook()
    ws = wb.active
    ws.title = "标书字段填写"

    headers = ["字段名称", "字段说明", "填写值", "出现章节"]
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    for row_idx, (field_name, chapter_titles) in enumerate(sorted(field_map.items()), 2):
        ws.cell(row=row_idx, column=1, value=field_name).border = thin_border
        ws.cell(row=row_idx, column=2, value="").border = thin_border
        ws.cell(row=row_idx, column=3, value="").border = thin_border
        ws.cell(row=row_idx, column=4, value=", ".join(chapter_titles)).border = thin_border

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 40
    ws.column_dimensions["D"].width = 40

    unlocked = Protection(locked=False)
    for row in range(2, len(field_map) + 2):
        ws.cell(row=row, column=2).protection = unlocked
        ws.cell(row=row, column=3).protection = unlocked

    Path("exports").mkdir(parents=True, exist_ok=True)
    suffix = f"_ch{chapter_index}" if chapter_index else ""
    filepath = f"exports/fields_{bid_id}{suffix}.xlsx"
    wb.save(filepath)

    return FileResponse(
        filepath,
        filename=f"{bid.name}_字段填写模板{suffix}.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@router.get("/{bid_id}", response_model=BidDetail)
def get_bid(bid_id: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")
    chapters = db.query(BidChapter).filter(
        BidChapter.bid_id == bid_id
    ).order_by(BidChapter.chapter_index).all()
    return BidDetail(
        id=bid.id, tender_id=bid.tender_id, name=bid.name,
        status=bid.status, create_time=bid.create_time,
        chapters=[BidChapterOut.model_validate(ch) for ch in chapters],
    )


@router.put("/{bid_id}/chapters/{chapter_index}")
def update_chapter(bid_id: int, chapter_index: int, content: dict, db: Session = Depends(get_db)):
    chapter = db.query(BidChapter).filter(
        BidChapter.bid_id == bid_id, BidChapter.chapter_index == chapter_index
    ).first()
    if not chapter:
        raise HTTPException(404, "章节不存在")

    chapter.content = content.get("content", chapter.content)
    chapter.last_modified = datetime.now()
    db.commit()
    return {"ok": True}


@router.post("/{bid_id}/export")
def export_bid(bid_id: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    from docx import Document

    doc = Document()
    doc.add_heading(bid.name, 0)

    chapters = db.query(BidChapter).filter(
        BidChapter.bid_id == bid_id
    ).order_by(BidChapter.chapter_index).all()

    for ch in chapters:
        doc.add_heading(ch.title, level=1)
        if ch.content:
            doc.add_paragraph(ch.content)

    Path("exports").mkdir(parents=True, exist_ok=True)
    filepath = f"exports/{bid.name}.docx"
    doc.save(filepath)

    return FileResponse(filepath, filename=f"{bid.name}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.post("/{bid_id}/fields/fill")
async def fill_fields(bid_id: int, file: UploadFile = File(...), chapter_index: int | None = Form(None), db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    if not file.filename or not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "仅支持 .xlsx 格式的 Excel 文件")

    import uuid
    temp_path = f"exports/fill_upload_{bid_id}_{uuid.uuid4().hex}.xlsx"
    content = await file.read()
    with open(temp_path, "wb") as f:
        f.write(content)

    from openpyxl import load_workbook
    wb = load_workbook(temp_path)
    ws = wb.active

    field_values: dict[str, str] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        field_name = str(row[0]).strip() if row[0] else ""
        fill_value = str(row[2]).strip() if len(row) > 2 and row[2] else ""
        if field_name and fill_value:
            field_values[field_name] = fill_value

    wb.close()
    try:
        os.remove(temp_path)
    except OSError:
        pass

    if not field_values:
        raise HTTPException(400, "未找到任何填写值，请确保Excel中第3列（填写值）有内容")

    query = db.query(BidChapter).filter(BidChapter.bid_id == bid_id)
    if chapter_index is not None:
        query = query.filter(BidChapter.chapter_index == chapter_index)
    chapters = query.all()

    updated_count = 0
    replaced_fields: set[str] = set()
    placeholder_pattern = re.compile(r'\{\{([^{}]+?)\}\}')

    for ch in chapters:
        if not ch.content:
            continue
        original = ch.content

        def replace_match(m):
            fname = m.group(1).strip()
            if fname in field_values:
                replaced_fields.add(fname)
                return field_values[fname]
            return m.group(0)

        new_content = placeholder_pattern.sub(replace_match, ch.content)
        if new_content != original:
            ch.content = new_content
            ch.last_modified = datetime.now()
            updated_count += 1

    db.commit()

    all_remaining: set[str] = set()
    for ch in chapters:
        if ch.content:
            for m in placeholder_pattern.finditer(ch.content):
                all_remaining.add(m.group(1).strip())

    return {
        "ok": True,
        "updated_chapters": updated_count,
        "filled_fields": sorted(replaced_fields),
        "unfilled_fields": sorted(all_remaining),
    }


@router.post("/{bid_id}/inspect", response_model=InspectResult)
async def inspect_bid_endpoint(bid_id: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")
    if not bid.tender_id:
        raise HTTPException(400, "该标书未关联招标文件，无法进行一键检查")

    from ..services.bid_inspector import inspect_bid as do_inspect
    missing_fields = await do_inspect(bid_id, db)

    cache_path = Path(f"exports/inspect_{bid_id}.json")
    cache_path.parent.mkdir(parents=True, exist_ok=True)
    with open(cache_path, "w", encoding="utf-8") as f:
        json.dump({"bid_id": bid_id, "missing_fields": missing_fields}, f, ensure_ascii=False)

    return InspectResult(
        bid_id=bid_id,
        missing_fields=[MissingField(**f) for f in missing_fields],
        total_count=len(missing_fields),
    )


@router.get("/{bid_id}/inspect/excel")
def export_inspect_excel(bid_id: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    cache_path = Path(f"exports/inspect_{bid_id}.json")
    if not cache_path.exists():
        raise HTTPException(400, "请先执行一键检查（POST /inspect）")

    with open(cache_path, "r", encoding="utf-8") as f:
        cached = json.load(f)

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, Protection

    wb = Workbook()
    ws = wb.active
    ws.title = "标书缺失信息"

    headers = ["字段名称", "字段说明", "填写值", "优先级", "所属章节", "类别"]
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border

    for row_idx, field in enumerate(cached["missing_fields"], 2):
        ws.cell(row=row_idx, column=1, value=field.get("field_name", "")).border = thin_border
        ws.cell(row=row_idx, column=2, value=field.get("description", "")).border = thin_border
        ws.cell(row=row_idx, column=3, value="").border = thin_border
        ws.cell(row=row_idx, column=4, value=field.get("priority", "")).border = thin_border
        ws.cell(row=row_idx, column=5, value=field.get("suggested_chapter_title", "")).border = thin_border
        ws.cell(row=row_idx, column=6, value=field.get("category", "")).border = thin_border

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 40
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 25
    ws.column_dimensions["F"].width = 10

    unlocked = Protection(locked=False)
    for row in range(2, len(cached["missing_fields"]) + 2):
        ws.cell(row=row, column=3).protection = unlocked

    filepath = f"exports/inspect_fields_{bid_id}.xlsx"
    wb.save(filepath)

    return FileResponse(
        filepath,
        filename=f"{bid.name}_缺失信息填写模板.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@router.post("/{bid_id}/inspect/fill")
async def fill_inspect_fields(bid_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")

    cache_path = Path(f"exports/inspect_{bid_id}.json")
    if not cache_path.exists():
        raise HTTPException(400, "请先执行一键检查（POST /inspect）")

    with open(cache_path, "r", encoding="utf-8") as f:
        cached = json.load(f)

    if not file.filename or not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "仅支持 .xlsx 格式的 Excel 文件")

    import uuid
    temp_path = f"exports/inspect_upload_{bid_id}_{uuid.uuid4().hex}.xlsx"
    content = await file.read()
    with open(temp_path, "wb") as f:
        f.write(content)

    from openpyxl import load_workbook
    wb = load_workbook(temp_path)
    ws = wb.active

    field_values: dict[str, str] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        field_name = str(row[0]).strip() if row[0] else ""
        fill_value = str(row[2]).strip() if len(row) > 2 and row[2] else ""
        if field_name and fill_value:
            field_values[field_name] = fill_value

    wb.close()
    try:
        os.remove(temp_path)
    except OSError:
        pass

    if not field_values:
        raise HTTPException(400, "未找到任何填写值，请确保Excel中第3列（填写值）有内容")

    from ..services.bid_inspector import fill_missing_fields_stream

    async def stream():
        async for event_text in fill_missing_fields_stream(
            bid_id, cached["missing_fields"], field_values, db
        ):
            yield event_text

    return StreamingResponse(stream(), media_type="text/event-stream")


@router.delete("/{bid_id}")
def delete_bid(bid_id: int, db: Session = Depends(get_db)):
    bid = db.query(Bid).filter(Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(404, "标书不存在")
    db.delete(bid)
    db.commit()
    return {"ok": True}
